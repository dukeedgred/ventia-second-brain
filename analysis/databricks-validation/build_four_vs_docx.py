from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("analysis/databricks-validation/transport-four-vs-boss-summary.docx")
FALLBACK_OUT = Path("analysis/databricks-validation/transport-four-vs-boss-summary-sht-corrected.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_paragraph(paragraph, size=11, bold=False, color=None, align=None, after=6):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.line_spacing = 1.1
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    style_paragraph(p, size=10.5)


def set_table_font(table, header_rows=1, body_size=8.5, header_size=8.5):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(header_size if r_idx < header_rows else body_size)
                    if r_idx < header_rows:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x29, 0x33, 0x46)
            if r_idx < header_rows:
                set_cell_shading(cell, "F2F4F7")


def add_table(doc, headers, rows, widths, body_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
    set_table_width(table, widths)
    set_table_font(table, body_size=body_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    p = doc.add_paragraph()
    title = p.add_run("Transport Four Vs Summary")
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.bold = True
    title.font.color.rgb = RGBColor(0x17, 0x20, 0x33)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.add_run("Generated: 09/06/2026").bold = True
    p.add_run(" | Scope: Asset Vision source tables and Transport contract/contractor schemas visible in Databricks")
    style_paragraph(p, size=10.5, color="5C667A", after=12)

    add_heading(doc, "Caveats", level=2)
    for item in [
        "Volume is row count by table object. Views and base tables can overlap, so totals are not deduplicated business records.",
        "Contract/contractor table volumes are accessible row counts. Some views could not be counted because Databricks returned permission errors or broken view dependencies.",
        "Velocity is inferred from the freshest timestamp/date values found in the tables. It is not the confirmed pipeline schedule.",
        "Usage/veracity means how much the data is used. Databricks lineage usage could not be measured because system.access.table_lineage returned USE SCHEMA permission denied.",
        "SHT/Maximo caveat: stakeholder documentation says Sydney Harbour Tunnel is operationally a Maximo tunnel contract. The Databricks source inventory also contains ext_mssql_asset_vision_vns_gen7 with a supplied source-system comment mapping it to SHT/WHT. That proves an Asset Vision-labelled Databricks source catalogue exists for SHT/WHT data; it does not prove SHT's authoritative operational system is Asset Vision.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Executive Summary", level=2)
    add_table(
        doc,
        ["Data area", "Volume", "Velocity", "Variety", "Usage / veracity"],
        [
            [
                "Asset Vision source tables",
                "111,802,224 rows across 230 table objects",
                "Current operational data present; latest source activity reaches 09/06/2026 or 10/06/2026 depending on timestamp/timezone handling",
                "Assets, jobs, inspections, defects/hazards, capital works, photos, forms/modules, workflow/status, resources, timesheets, reference/export/system data",
                "Actual query/lineage usage not available with current permissions",
            ],
            [
                "Transport contract/contractor tables",
                "233,560,284 accessible rows across 319 table objects; 122 objects counted and 197 objects not executable/countable from current auth",
                "Mixed: most active schemas have current data; RAMC/SRAP-C are recent; SHT latest counted timestamp is 14/04/2026",
                "Traffic counts, KPI/reporting views, job exports, timesheets, capital works, lane access, tunnel/contract-specific outputs, reference/date/EOT tables",
                "Actual query/lineage usage not available with current permissions",
            ],
        ],
        [1700, 1800, 2200, 2500, 2100],
        body_size=8,
    )

    add_heading(doc, "Source Table Summary", level=2)
    add_table(
        doc,
        ["Source catalogue", "Objects", "Row count", "Velocity signal", "Variety"],
        [
            ["ext_mssql_asset_vision_ven_gen7", "38", "24,051,436", "Current source activity present", "Asset Vision operational core: assets, jobs, inspections, capital works, photos, forms/modules, workflow, resources, timesheets"],
            ["ext_mssql_asset_vision_ven_rms", "38", "24,569,468", "Current source activity present", "Asset Vision operational core, SRAP-C/RMS context"],
            ["ext_mssql_asset_vision_ven_vicroads", "40", "37,854,148", "Current source activity present", "Asset Vision operational core plus lane access and SQL Server version metadata"],
            ["ext_mssql_asset_vision_vns_gen7", "38", "5,071,032", "Current source activity present", "Asset Vision-labelled Databricks source catalogue mapped to SHT/WHT; operational SHT system remains documented as Maximo"],
            ["ext_mssql_asset_vision_vnz_gen7", "38", "14,430,838", "Current source activity present", "Asset Vision operational core, Auckland West context"],
            ["ext_mssql_asset_vision_vsm_gen7", "38", "5,825,302", "Current source activity present", "Asset Vision operational core, VentureSmart context"],
        ],
        [2600, 900, 1300, 1800, 3600],
        body_size=8.5,
    )

    add_heading(doc, "Contract / Contractor Table Summary", level=2)
    add_table(
        doc,
        ["Contract schema", "Objects", "Counted", "Accessible rows", "Velocity", "Count gaps"],
        [
            ["transport_dev.transport", "40", "8", "16,519,404", "Current: within 1 day", "32 views not counted due permission errors"],
            ["transport_dev.transport_aklw", "11", "11", "1,822,984", "Current: within 1 day", "None from current run"],
            ["transport_dev.transport_fndc", "12", "12", "264,594", "Current: within 1 day", "None from current run"],
            ["transport_dev.transport_nel", "4", "4", "11,405", "Current: within 1 day", "None from current run"],
            ["transport_dev.transport_ramc", "17", "9", "39,868", "Recent: within 31 days", "8 views not counted due permission errors"],
            ["transport_dev.transport_sht", "27", "13", "165,474,873", "Stale: latest counted timestamp 14/04/2026", "14 views not counted due unresolved column/view issues"],
            ["transport_dev.transport_srapc", "42", "5", "1,108", "Recent: within 31 days", "37 views not counted due broken dependencies"],
            ["transport_dev.transport_tsrc", "88", "42", "829,475", "Current: within 1 day", "46 views not counted due permission errors or invalid view SQL"],
            ["transport_dev.transport_wru", "78", "18", "48,596,573", "Current: within 1 day", "60 views not counted due permission errors"],
        ],
        [2450, 700, 750, 1400, 1900, 2850],
        body_size=8,
    )

    add_heading(doc, "What Each Contract / Contractor Schema Contains", level=2)
    add_table(
        doc,
        ["Contract schema", "Sort of data present", "Table/view signals"],
        [
            [
                "transport_dev.transport",
                "Cross-contract transport commercial, workforce and work-order support data.",
                "All-contract jobs, job form fields, resources, timesheets, SAP items, purchase orders, purchase requisitions, vendor master/open/cleared items, cost centre/WBS/project master data, equipment master/depreciation, service entry and CATS/PTMW views.",
            ],
            [
                "transport_dev.transport_aklw",
                "Auckland West operational Asset Vision-style data.",
                "Asset, job, inspection, capital work, capital work task, form fields, workflow status, timesheet item/report, plant pending timesheet and updated dispatch ID views.",
            ],
            [
                "transport_dev.transport_fndc",
                "Road asset layers, forward works, dispatch/cost and weather data.",
                "Surface, drainage, footpath, pavement layer, railings, signs, treatment length, forward work view, maintenance cost, dispatch, dispatch claim and hourly weather forecast tables/views.",
            ],
            [
                "transport_dev.transport_nel",
                "KPI-oriented asset, work-order and system availability data.",
                "KPI assets, KPI work orders, KPI system availability devices and reference date table.",
            ],
            [
                "transport_dev.transport_ramc",
                "Backlog, monthly job snapshot, strip-map and job/inspection reporting data.",
                "Backlog change reports, current/last month job snapshots, monthly backlog reduction, reporting period, job, inspection, road-last-inspected and strip-map views.",
            ],
            [
                "transport_dev.transport_sht",
                "SHT contract reporting/curated data. This is not the same thing as saying the source operational system is Asset Vision.",
                "NPS/SPS/TVS northbound/southbound segment tables, segmented views, all assets, critical assets, job, inspection, user groups, AI views and North Sydney rolling weather view. Stakeholder notes identify SHT operational work management as Maximo.",
            ],
            [
                "transport_dev.transport_srapc",
                "SRAP-C asset, job, inspection, customer request, monthly report and subcontractor reporting data.",
                "Monthly report, Formitize mapping, TACP constants/TOC/data loads, civil master, bridge/slope/school-zone assets, asset/job/inspection/customer request attribute views, photos, monthly subcontractor energy/health-safety/material/social/fuel/waste views, special projects, TFNSW monthly defect intervention and weather observations.",
            ],
            [
                "transport_dev.transport_tsrc",
                "TSRC assets, incidents, work orders, inspections, KPI/abatement, PCAS/pavement, traffic and lane-closure data.",
                "AED assets/incidents/work orders, asset register, work order, inspection requirements, road safety audit, KPI 2/3/10/18/19/25 views, ITS asset uptime/jobs, CCTV requests, toll road unavailability, PCAS/pavement tables, lane closure reference factors, traffic closures/volume, customer requests and capital work views.",
            ],
            [
                "transport_dev.transport_wru",
                "WRU traffic counts, lane access, road asset categories, inspections, jobs, capital works, documents and timesheets.",
                "Counter locations, counts by carriageway/lane/hour, lane access config/chainage/traffic volumes, asset category views, bridge/BIS/RAS views, inspection dashboards/levels/hazard-defect/completions, job/job geometry/job export, capital work, EOT, documents, timesheets, TGS speed limits, public holidays and work peak periods.",
            ],
        ],
        [2100, 2800, 4460],
        body_size=7.6,
    )

    add_heading(doc, "Plain-Language Takeaways", level=2)
    for item in [
        "Volume: the data estate is large. Source tables hold about 111.8M table-object rows. Contract/contractor schemas hold at least 233.6M accessible rows, with more behind views that could not be executed with the current account.",
        "Velocity: this is not a static archive. Current timestamp signals exist in the main source tables and most contract schemas. SHT is the main visible stale area in the counted contract tables.",
        "Variety: the source layer is operational Asset Vision data. The contract layer is more reporting/curation-heavy, especially traffic counts, KPI outputs, timesheets, lane access, capital works and contract-specific views.",
        "Usage/veracity: actual usage counts were not available because Databricks system lineage access is blocked. Claims about most-used tables require system.access.table_lineage, query history, dashboard lineage, or BI/report access logs.",
        "SHT system interpretation: treat SHT as Maximo-led operationally unless the tunnel SMEs or source owners confirm otherwise. Treat ext_mssql_asset_vision_vns_gen7 as an Asset Vision-labelled Databricks source catalogue that contains SHT/WHT-mapped data, not as proof that SHT operates in Asset Vision.",
    ]:
        add_bullet(doc, item)

    footer = section.footer.paragraphs[0]
    footer.text = "Transport Four Vs Summary | Generated 09/06/2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(footer, size=8.5, color="5C667A", after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT)
        print(OUT)
    except PermissionError:
        doc.save(FALLBACK_OUT)
        print(FALLBACK_OUT)


if __name__ == "__main__":
    main()
